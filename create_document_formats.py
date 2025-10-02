#!/usr/bin/env python3
"""
Create document formats for Google Ads API Design Document
"""

import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import markdown

def create_docx():
    """Create DOCX version"""
    try:
        # Read markdown content
        with open('GOOGLE_ADS_API_DESIGN_DOCUMENT.md', 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Marketing ROI Tool - Google Ads API Integration Design Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add document info
        doc.add_paragraph('Document Version: 1.0')
        doc.add_paragraph('Date: January 2025')
        doc.add_paragraph('Company: SiteOptz.ai')
        doc.add_paragraph('Contact: [Your Contact Information]')
        doc.add_paragraph('')
        
        # Process content
        lines = content.split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                doc.add_paragraph('')
                continue
                
            # Handle headers
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('- '):
                # Bullet point
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                p = doc.add_paragraph(line[2:-2])
                p.runs[0].bold = True
            elif line.startswith('```'):
                # Code block
                continue  # Skip code block markers for now
            else:
                # Regular paragraph
                doc.add_paragraph(line)
        
        # Save document
        doc.save('Google_Ads_API_Design_Document.docx')
        print("✅ DOCX created: Google_Ads_API_Design_Document.docx")
        
    except Exception as e:
        print(f"❌ Error creating DOCX: {e}")

def create_rtf():
    """Create RTF version"""
    try:
        # Read markdown content
        with open('GOOGLE_ADS_API_DESIGN_DOCUMENT.md', 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Convert to HTML first
        html = markdown.markdown(content)
        
        # Basic RTF conversion
        rtf_content = f"""{{\\rtf1\\ansi\\deff0 {{
\\fonttbl {{\\f0 Times New Roman;}}
\\colortbl;\\red0\\green0\\blue0;
\\margl1440\\margr1440\\margt1440\\margb1440

\\f0\\fs24
"""
        
        # Simple HTML to RTF conversion
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            
            if line.startswith('# '):
                rtf_content += f"\\b {line[2:]}\\b0\\par\\par\n"
            elif line.startswith('## '):
                rtf_content += f"\\b {line[3:]}\\b0\\par\\par\n"
            elif line.startswith('### '):
                rtf_content += f"\\b {line[4:]}\\b0\\par\n"
            elif line.startswith('- '):
                rtf_content += f"\\bullet {line[2:]}\\par\n"
            elif line.startswith('**') and line.endswith('**'):
                rtf_content += f"\\b {line[2:-2]}\\b0\\par\n"
            elif line:
                rtf_content += f"{line}\\par\n"
            else:
                rtf_content += "\\par\n"
        
        rtf_content += "}"
        
        # Save RTF file
        with open('Google_Ads_API_Design_Document.rtf', 'w', encoding='utf-8') as f:
            f.write(rtf_content)
        
        print("✅ RTF created: Google_Ads_API_Design_Document.rtf")
        
    except Exception as e:
        print(f"❌ Error creating RTF: {e}")

def create_simple_pdf():
    """Create a simple PDF version using basic formatting"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        # Read markdown content
        with open('GOOGLE_ADS_API_DESIGN_DOCUMENT.md', 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Create PDF
        doc = SimpleDocTemplate("Google_Ads_API_Design_Document.pdf", pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=8
        )
        
        # Build content
        story = []
        
        # Title
        story.append(Paragraph("Marketing ROI Tool - Google Ads API Integration Design Document", title_style))
        story.append(Spacer(1, 12))
        
        # Document info
        story.append(Paragraph("Document Version: 1.0", styles['Normal']))
        story.append(Paragraph("Date: January 2025", styles['Normal']))
        story.append(Paragraph("Company: SiteOptz.ai", styles['Normal']))
        story.append(Paragraph("Contact: [Your Contact Information]", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Process content
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            
            if not line:
                story.append(Spacer(1, 6))
                continue
                
            if line.startswith('# '):
                story.append(Paragraph(line[2:], title_style))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], heading_style))
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], styles['Heading3']))
            elif line.startswith('- '):
                story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
            elif line.startswith('**') and line.endswith('**'):
                story.append(Paragraph(f"<b>{line[2:-2]}</b>", styles['Normal']))
            else:
                story.append(Paragraph(line, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        print("✅ PDF created: Google_Ads_API_Design_Document.pdf")
        
    except Exception as e:
        print(f"❌ Error creating PDF: {e}")

def main():
    """Main function"""
    print("🚀 Creating document formats...")
    print("=" * 50)
    
    if not os.path.exists('GOOGLE_ADS_API_DESIGN_DOCUMENT.md'):
        print("❌ Source file not found: GOOGLE_ADS_API_DESIGN_DOCUMENT.md")
        return
    
    # Create HTML version
    try:
        with open('GOOGLE_ADS_API_DESIGN_DOCUMENT.md', 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        html = markdown.markdown(md_content, extensions=['toc', 'tables', 'codehilite'])
        
        full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Google Ads API Design Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        h1 {{ color: #1a73e8; border-bottom: 2px solid #1a73e8; padding-bottom: 10px; }}
        h2 {{ color: #34a853; margin-top: 30px; }}
        h3 {{ color: #ea4335; }}
        code {{ background-color: #f8f9fa; padding: 2px 4px; border-radius: 3px; }}
        pre {{ background-color: #f8f9fa; padding: 15px; border-radius: 5px; overflow-x: auto; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
    </style>
</head>
<body>
    {html}
</body>
</html>"""
        
        with open('Google_Ads_API_Design_Document.html', 'w', encoding='utf-8') as f:
            f.write(full_html)
        
        print("✅ HTML created: Google_Ads_API_Design_Document.html")
        
    except Exception as e:
        print(f"❌ Error creating HTML: {e}")
    
    # Create other formats
    create_docx()
    create_rtf()
    create_simple_pdf()
    
    print("=" * 50)
    print("✅ Document creation complete!")
    
    # List created files
    files = [
        'Google_Ads_API_Design_Document.html',
        'Google_Ads_API_Design_Document.pdf',
        'Google_Ads_API_Design_Document.docx',
        'Google_Ads_API_Design_Document.rtf'
    ]
    
    print("\nGenerated files:")
    for file in files:
        if os.path.exists(file):
            size = os.path.getsize(file) / 1024
            print(f"  📄 {file} ({size:.1f} KB)")

if __name__ == "__main__":
    main()
